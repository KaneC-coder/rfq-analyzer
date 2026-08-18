"""
RFQ智能报价助手 - Web应用 (Vercel Serverless 版本)
"""

import os
import json
import uuid
import threading
from datetime import datetime

from flask import Flask, request, jsonify, render_template, send_from_directory

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'rfq-analyzer-secret-key-2026')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

BASE_DIR = '/tmp/rfq_data'
UPLOAD_DIR = os.path.join(BASE_DIR, 'uploads')
OUTPUT_DIR = os.path.join(BASE_DIR, 'outputs')
for d in [UPLOAD_DIR, OUTPUT_DIR]:
    os.makedirs(d, exist_ok=True)

USERS_FILE = os.path.join(BASE_DIR, 'users.json')
RECORDS_FILE = os.path.join(BASE_DIR, 'records.json')

_lock = threading.Lock()

def _load_json(filepath, default):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default

def _save_json(filepath, data):
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_or_create_user(ip):
    with _lock:
        users = _load_json(USERS_FILE, {})
        if ip not in users:
            users[ip] = {'ip': ip, 'free_count': 3, 'total_used': 0,
                         'created_at': datetime.utcnow().isoformat(),
                         'last_used_at': datetime.utcnow().isoformat()}
            _save_json(USERS_FILE, users)
        return users[ip]

def update_user(ip, free_count=None, total_used=None):
    with _lock:
        users = _load_json(USERS_FILE, {})
        if ip in users:
            if free_count is not None:
                users[ip]['free_count'] = free_count
            if total_used is not None:
                users[ip]['total_used'] = total_used
            users[ip]['last_used_at'] = datetime.utcnow().isoformat()
            _save_json(USERS_FILE, users)
            return users[ip]
        return None

def create_record(user_ip, image_path):
    with _lock:
        records = _load_json(RECORDS_FILE, {})
        rid = str(uuid.uuid4().hex[:12])
        records[rid] = {'id': rid, 'user_ip': user_ip, 'image_path': image_path,
                        'product_info': None, 'report_path': None,
                        'status': 'processing', 'created_at': datetime.utcnow().isoformat()}
        _save_json(RECORDS_FILE, records)
        return records[rid]

def get_record(rid):
    return _load_json(RECORDS_FILE, {}).get(rid)

def update_record(rid, **kwargs):
    with _lock:
        records = _load_json(RECORDS_FILE, {})
        if rid in records:
            records[rid].update(kwargs)
            _save_json(RECORDS_FILE, records)
            return records[rid]
        return None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/init')
def init_session():
    ip = request.remote_addr
    user = get_or_create_user(ip)
    return jsonify({'success': True, 'free_count': user['free_count'], 'total_used': user['total_used']})

@app.route('/api/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    if not file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.webp')):
        return jsonify({'error': '只支持JPG、PNG、WEBP格式'}), 400

    ip = request.remote_addr
    user = get_or_create_user(ip)
    if user['free_count'] <= 0:
        return jsonify({'error': '免费次数已用完，请付费继续使用（0.5元/次）', 'need_payment': True}), 402

    ext = os.path.splitext(file.filename)[1].lower()
    filename = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(UPLOAD_DIR, filename)
    file.save(filepath)

    record = create_record(ip, filepath)
    return jsonify({'success': True, 'record_id': record['id'],
                    'free_count': user['free_count'],
                    'message': f"上传成功，剩余免费次数：{user['free_count']}次"})

@app.route('/api/analyze/<record_id>', methods=['POST'])
def analyze(record_id):
    record = get_record(record_id)
    if not record:
        return jsonify({'error': '记录不存在'}), 404
    ip = request.remote_addr
    user = get_or_create_user(ip)
    if user['free_count'] <= 0:
        return jsonify({'error': '免费次数已用完', 'need_payment': True}), 402

    try:
        product_info = {
            'product_name': '识别到的产品名称（待接入OCR）',
            'quantity': '数量',
            'buyer_country': '买家国家',
            'specifications': '规格参数'
        }
        update_record(record_id, product_info=json.dumps(product_info, ensure_ascii=False), status='completed')
        new_free = user['free_count'] - 1
        new_total = user['total_used'] + 1
        update_user(ip, free_count=new_free, total_used=new_total)
        return jsonify({'success': True, 'product_info': product_info, 'free_count': new_free})
    except Exception as e:
        update_record(record_id, status='failed')
        return jsonify({'error': str(e)}), 500

@app.route('/api/report/<record_id>', methods=['POST'])
def generate_report(record_id):
    record = get_record(record_id)
    if not record or record.get('status') != 'completed':
        return jsonify({'error': '分析未完成'}), 400
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        product_info = json.loads(record['product_info']) if record.get('product_info') else {}
        doc = Document()
        title = doc.add_heading('RFQ报价分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_heading('1. 买家信息', level=1)
        doc.add_paragraph(f"买家国家：{product_info.get('buyer_country', '未知')}")
        doc.add_heading('2. 产品需求', level=1)
        doc.add_paragraph(f"产品名称：{product_info.get('product_name', '未知')}")
        doc.add_paragraph(f"数量：{product_info.get('quantity', '未知')}")
        doc.add_paragraph(f"规格：{product_info.get('specifications', '未知')}")
        doc.add_heading('3. 价格与供应商参考', level=1)
        doc.add_paragraph('（待接入真实搜索数据）')
        doc.add_heading('4. 关税与报关信息', level=1)
        doc.add_paragraph('（待接入真实关税数据）')
        doc.add_heading('5. 利润评估', level=1)
        doc.add_paragraph('（待接入真实成本计算）')
        doc.add_heading('6. 开发信模板', level=1)
        doc.add_paragraph('（待生成）')
        doc.add_heading('7. 风险提示', level=1)
        doc.add_paragraph('以上分析仅供参考，请结合实际市场情况做出决策。')
        report_filename = f"RFQ_Report_{record_id}.docx"
        report_path = os.path.join(OUTPUT_DIR, report_filename)
        doc.save(report_path)
        update_record(record_id, report_path=report_path)
        return jsonify({'success': True, 'report_path': report_path, 'report_filename': report_filename})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<record_id>')
def download_report(record_id):
    record = get_record(record_id)
    if not record or not record.get('report_path'):
        return jsonify({'error': '报告未生成'}), 404
    return send_from_directory(OUTPUT_DIR, os.path.basename(record['report_path']), as_attachment=True)

@app.route('/api/status/<record_id>')
def get_status(record_id):
    record = get_record(record_id)
    if not record:
        return jsonify({'error': '记录不存在'}), 404
    ip = request.remote_addr
    user = get_or_create_user(ip)
    product_info = None
    if record.get('product_info'):
        try:
            product_info = json.loads(record['product_info'])
        except:
            pass
    return jsonify({'record_id': record_id, 'status': record.get('status'),
                    'free_count': user['free_count'], 'product_info': product_info})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
